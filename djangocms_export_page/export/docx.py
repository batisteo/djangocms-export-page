
from io import BytesIO
from os import path
from django.utils.encoding import force_text

from docx import Document
from docx.shared import Cm

# from bureaucracy import DocxTemplate, HTML, Table
# from bureaucracy.replacements import ListOfReplacement

package_path = path.abspath(path.dirname(path.dirname(__file__)))

DOCX_PATH = path.join(package_path, 'templates', 'djangocms_export_page', 'template.docx')


def export_to_docx(data):
    document = Document()

    for placeholder, plugins in data.items():
        heading = placeholder.slot.replace('-', ' ').title()
        document.add_heading(heading, level=3)

        for plugin, fields in plugins.items():
            rows = []

            for field_name, field in fields.items():
                value = getattr(plugin, field_name)
                if value:
                    rows.append([field.verbose_name, value])

            if rows:
                heading = ': '.join((force_text(s) for s in [plugin._meta.model.__name__, plugin]))
                document.add_heading(heading, level=5)
                add_table(document, rows, has_header=False)

    return save(document)


def save(document):
    handle = BytesIO()
    document.save(handle)
    handle.seek(0)
    return handle.read()


def add_table(document, rows, has_header=False):
    table = document.add_table(rows=0, cols=len(rows[0]))

    if has_header:
        populate_row(table, rows.pop(0))

    for row in rows:
        populate_row(table, row)

    table.style = 'TableGrid'
    return table


def populate_row(table, row):
    cells = table.add_row().cells
    for i, content in enumerate(row):
        cells[i].text = str(content)
